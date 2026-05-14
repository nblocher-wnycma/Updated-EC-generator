"""Word export for Existing Conditions narratives."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .file_manager import combined_docx_path, section_docx_path
from .input_prompts import get_section_title
from .quality_control import QualityIssue


def export_section_docx(
    general: dict[str, str],
    section_key: str,
    narrative: str,
    issues: list[QualityIssue],
    output_path: Path | None = None,
) -> Path:
    path = output_path or section_docx_path(general, section_key)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = _base_document(general, f"Existing Conditions - {get_section_title(section_key)}")
    _add_narrative(doc, get_section_title(section_key), narrative)
    _add_qc_summary(doc, issues)
    doc.save(path)
    return path


def export_combined_docx(
    general: dict[str, str],
    narratives: dict[str, str],
    issues_by_section: dict[str, list[QualityIssue]],
    output_path: Path | None = None,
) -> Path:
    path = output_path or combined_docx_path(general)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = _base_document(general, "Existing Conditions Narratives")
    for index, (section_key, narrative) in enumerate(narratives.items()):
        if index:
            doc.add_page_break()
        _add_narrative(doc, get_section_title(section_key), narrative)
        _add_qc_summary(doc, issues_by_section.get(section_key, []))
    doc.save(path)
    return path


def _base_document(general: dict[str, str], title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(31, 78, 121)

    farm = doc.add_paragraph()
    farm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    farm_run = farm.add_run(general.get("farm_name", "Unnamed Farm"))
    farm_run.bold = True
    farm_run.font.size = Pt(13)

    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for label, key in [
        ("Owner/Operator", "owner_operator"),
        ("Address", "farm_address"),
        ("County/State", "county_state"),
        ("Operation Type", "operation_type"),
        ("Evaluation Date", "evaluation_date"),
        ("Evaluator", "evaluator"),
    ]:
        value = general.get(key, "")
        if value:
            row = meta.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
    doc.add_paragraph()
    return doc


def _add_narrative(doc: Document, section_title: str, narrative: str) -> None:
    doc.add_heading(section_title, level=1)
    for paragraph_text in narrative.split("\n\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.08
        paragraph.add_run(paragraph_text.strip())


def _add_qc_summary(doc: Document, issues: list[QualityIssue]) -> None:
    doc.add_paragraph()
    doc.add_heading("Quality-Control Notes", level=2)
    if not issues:
        doc.add_paragraph("No quality-control warnings were identified at export.")
        return
    for issue in issues:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{issue.severity.upper()}: ").bold = True
        paragraph.add_run(issue.message)
